import tkinter as tk
from tkinter import ttk, messagebox
from docx import Document


def generate_word():
    document = Document()
    document.add_heading('Biodata', 0)

    # Personal Information
    document.add_heading('Personal Information', level=1)
    document.add_paragraph(f"First Name: {le_1.get('1.0', 'end-1c')}")
    document.add_paragraph(f"Middle Name: {le_2.get('1.0', 'end-1c')}")
    document.add_paragraph(f"Last Name: {le_3.get('1.0', 'end-1c')}")
    document.add_paragraph(f"Date of Birth: {le_4.get()}")

    # Contact Information
    document.add_heading('Contact Information', level=1)
    document.add_paragraph(f"Phone Number: {le_5.get('1.0', 'end-1c')}")
    document.add_paragraph(f"Email Address: {le_6.get()}")
    document.add_paragraph(f"Address: {le_7.get('1.0', 'end-1c')}")

    # Objective
    document.add_heading('Objective', level=1)
    document.add_paragraph(le_8.get())

    # Educational Background
    document.add_heading('Educational Background', level=1)
    document.add_paragraph(le_9.get())

    # Work Experience
    document.add_heading('Work Experience', level=1)
    document.add_paragraph(le_10.get())

    # Skills
    document.add_heading('Skills', level=1)
    document.add_paragraph(le_11.get())

    # Projects
    document.add_heading('Projects', level=1)
    document.add_paragraph(le_12.get())

    # Certifications
    document.add_heading('Certifications', level=1)
    document.add_paragraph(le_13.get())

    # Awards & Achievements
    document.add_heading('Awards & Achievements', level=1)
    document.add_paragraph(le_14.get())

    # Languages Known
    document.add_heading('Languages Known', level=1)
    document.add_paragraph(le_15.get())

    # Hobbies and Interests
    document.add_heading('Hobbies and Interests', level=1)
    document.add_paragraph(le_16.get())

    # References
    document.add_heading('References', level=1)
    document.add_paragraph(le_17.get())

    document.save('biodata.docx')
    messagebox.showinfo("Success", "Word file generated successfully!")


# Set up the main application window
root = tk.Tk()
root.title("Biodata Builder")
root.geometry("401x630")
root.configure(background="gray")


def clear_placeholder(event):
    if le_4.get() == ("DD/MM/YY"):
        le_4.delete(0, tk.END)


def clear_placeholder2(event):
    if le_8.get() == ("Type here a brief statement about your career goals and what you aim to achieve"):
        le_8.delete(0, tk.END)


def clear_placeholder3(event):
    if le_9.get() == ("List your educational qualifications here in reverse chronological order (most recent first)"):
        le_9.delete(0, tk.END)


def clear_placeholder4(event):
    if le_10.get() == ("List your work experience here in reverse chronological order"):
        le_10.delete(0, tk.END)


def clear_placeholder5(event):
    if le_11.get() == ("List your relevant technical and soft skills here."):
        le_11.delete(0, tk.END)


def clear_placeholder6(event):
    if le_12.get() == ("Type Brief descriptions of significant projects you have worked on."):
        le_12.delete(0, tk.END)


def clear_placeholder7(event):
    if le_13.get() == ("Type Any relevant certifications you have obtained."):
        le_13.delete(0, tk.END)


def clear_placeholder8(event):
    if le_14.get() == ("Type Any awards or significant achievements related to your career or education"):
        le_14.delete(0, tk.END)


def clear_placeholder9(event):
    if le_15.get() == ("List the languages you speak/read/write."):
        le_15.delete(0, tk.END)


def clear_placeholder10(event):
    if le_16.get() == ("A brief list of your hobbies and interests."):
        le_16.delete(0, tk.END)


def clear_placeholder11(event):
    if le_17.get() == ("Type references if requested, or mention that they are available upon request"):
        le_17.delete(0, tk.END)


# Personal Information
l = ttk.Label(root, text="...Personal Information...", font="Arial, 10", foreground="green", background="black")
l.grid(row=0, columnspan=10, sticky='nsew', padx=1, pady=7)
# First Name
l1 = ttk.Label(root, text="First Name", font="Arial, 7")
l1.grid(row=1, column=0, sticky=tk.W + tk.E, padx=1, pady=2)
le_1 = tk.Text(root, height=1, width=10, font="Arial, 8")
le_1.grid(row=1, column=1, sticky=tk.W + tk.E, padx=1, pady=2)
# Middle Name
l2 = ttk.Label(root, text="Middle Name", font="Arial, 7")
l2.grid(row=1, column=2, sticky=tk.W + tk.E, padx=1, pady=2)
le_2 = tk.Text(root, height=1, width=10, font="Arial, 8")
le_2.grid(row=1, column=3, sticky=tk.W + tk.E, padx=1, pady=2)
# Last Name
l3 = ttk.Label(root, text="Last Name", font="Arial, 7")
l3.grid(row=1, column=4, sticky=tk.W + tk.E, padx=1, pady=2)
le_3 = tk.Text(root, height=1, width=10, font="Arial, 8")
le_3.grid(row=1, column=5, sticky=tk.W + tk.E, pady=2, padx=1)
# Date of Birth
l4 = ttk.Label(root, text="Date of Birth", font="Arial, 7")
l4.grid(row=2, column=0, sticky=tk.W + tk.E, padx=1, pady=2)
le_4 = tk.Entry(root, width=10, font="Arial, 8")
le_4.insert(0, "DD/MM/YY")
le_4.bind("<FocusIn>", clear_placeholder)
le_4.grid(row=2, column=1, sticky='nsew', pady=2, padx=1)
# Contact Informations
# Phone Number
le5 = ttk.Label(root, text="Phone Number", font="Arial, 7")
le5.grid(row=2, column=2, sticky='nsew', pady=2, padx=1)
le_5 = tk.Text(root, height=1, width=10, font="Arial, 7")
le_5.grid(row=2, column=3, sticky='nsew', pady=2, padx=1)
# Email Address
le6 = ttk.Label(root, text="Email Address", font="Arial, 7")
le6.grid(row=2, column=4, sticky='nsew', padx=1, pady=2)
le_6 = tk.Entry(root, width=10, font="Arial, 8")
le_6.insert(0, "@gmail.com")
le_6.grid(row=2, column=5, sticky='nsew', padx=1, pady=2)
# Address
le7 = ttk.Label(root, text="Full Address", font="Arial, 7")
le7.grid(row=3, columnspan=1, sticky='nw', padx=1, pady=2)
le_7 = tk.Text(root, height=2, width=40, font="Arial, 8")
le_7.grid(row=4, columnspan=12, sticky='nsew', padx=1, pady=2)
# Objective
le8 = ttk.Label(root, text="Objective", font="Arial, 10", foreground="green", background="black")
le8.grid(row=5, columnspan=10, sticky='nsew', padx=1, pady=2)
le_8 = tk.Entry(root, width=40, font="Arial, 8")
le_8.insert(0, "Type here a brief statement about your career goals and what you aim to achieve")
le_8.bind("<FocusIn>", clear_placeholder2)
le_8.grid(row=6, columnspan=12, sticky='nsew', padx=1, pady=2)
# Educational Background
le9 = ttk.Label(root, text="Educational Background", font="Arial, 10", foreground="green", background="black")
le9.grid(row=7, columnspan=10, sticky='nsew', padx=1, pady=2)
le_9 = tk.Entry(root, width=40, font="Arial, 8")
le_9.insert(0, "List your educational qualifications here in reverse chronological order (most recent first)")
le_9.bind("<FocusIn>", clear_placeholder3)
le_9.grid(row=8, columnspan=12, sticky='nsew', padx=1, pady=2)
# Work Experience
le10 = ttk.Label(root, text="Work Experience", font="Arial, 10", foreground="green", background="black")
le10.grid(row=9, columnspan=10, sticky='nsew', padx=1, pady=2)
le_10 = tk.Entry(root, width=40, font="Arial, 8")
le_10.insert(0, "List your work experience here in reverse chronological order")
le_10.bind("<FocusIn>", clear_placeholder4)
le_10.grid(row=10, columnspan=12, sticky='nsew', padx=1, pady=2)
# Skills
le11 = ttk.Label(root, text="Skills", font="Arial, 10", foreground="green", background="black")
le11.grid(row=11, columnspan=10, sticky='nsew', padx=1, pady=2)
le_11 = tk.Entry(root, width=40, font="Arial, 8")
le_11.insert(0, "List your relevant technical and soft skills here.")
le_11.bind("<FocusIn>", clear_placeholder5)
le_11.grid(row=12, columnspan=12, sticky='nsew', padx=1, pady=2)
# Projects
le12 = ttk.Label(root, text="Projects", font="Arial, 10", foreground="green", background="black")
le12.grid(row=13, columnspan=10, sticky='nsew', padx=1, pady=2)
le_12 = tk.Entry(root, width=40, font="Arial, 8")
le_12.insert(0, "Type Brief descriptions of significant projects you have worked on.")
le_12.bind("<FocusIn>", clear_placeholder6)
le_12.grid(row=14, columnspan=12, sticky='nsew', padx=1, pady=2)
# Certifications
le13 = ttk.Label(root, text="Certifications", font="Arial, 10", foreground="green", background="black")
le13.grid(row=15, columnspan=10, sticky='nsew', padx=1, pady=2)
le_13 = tk.Entry(root, width=40, font="Arial, 8")
le_13.insert(0, "Type Any relevant certifications you have obtained.")
le_13.bind("<FocusIn>", clear_placeholder7)
le_13.grid(row=16, columnspan=12, sticky='nsew', padx=1, pady=2)
# Awards & Achievements
le14 = ttk.Label(root, text="Awards & Achievements", font="Arial, 10", foreground="green", background="black")
le14.grid(row=17, columnspan=10, sticky='nsew', padx=1, pady=2)
le_14 = tk.Entry(root, width=40, font="Arial, 8")
le_14.insert(0, "Type Any awards or significant achievements related to your career or education")
le_14.bind("<FocusIn>", clear_placeholder8)
le_14.grid(row=18, columnspan=12, sticky='nsew', padx=1, pady=2)
# Languages Known
le15 = ttk.Label(root, text="Languages Known", font="Arial, 10", foreground="green", background="black")
le15.grid(row=19, columnspan=10, sticky='nsew', padx=1, pady=2)
le_15 = tk.Entry(root, width=40, font="Arial, 8")
le_15.insert(0, "List the languages you speak/read/write.")
le_15.bind("<FocusIn>", clear_placeholder9)
le_15.grid(row=20, columnspan=12, sticky='nsew', padx=1, pady=2)
# Hobbies and Interests
le16 = ttk.Label(root, text="Hobbies and Interests", font="Arial, 10", foreground="green", background="black")
le16.grid(row=21, columnspan=10, sticky='nsew', padx=1, pady=2)
le_16 = tk.Entry(root, width=40, font="Arial, 8")
le_16.insert(0, "A brief list of your hobbies and interests.")
le_16.bind("<FocusIn>", clear_placeholder10)
le_16.grid(row=22, columnspan=12, sticky='nsew', padx=1, pady=2)
# References
le17 = ttk.Label(root, text="References", font="Arial, 10", foreground="green", background="black")
le17.grid(row=23, columnspan=10, sticky='nsew', padx=1, pady=2)
le_17 = tk.Entry(root, width=40, font="Arial, 8")
le_17.insert(0, "Type references if requested, or mention that they are available upon request")
le_17.bind("<FocusIn>", clear_placeholder11)
le_17.grid(row=24, columnspan=12, sticky='nsew', padx=1, pady=2)
# Generate Word File Button
generate_button = ttk.Button(root, text="Generate Word File", command=generate_word)
generate_button.grid(row=25, columnspan=12, pady=10)

root.mainloop()
